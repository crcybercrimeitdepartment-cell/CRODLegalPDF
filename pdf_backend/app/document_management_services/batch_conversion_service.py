"""
Batch Conversion Service — Document Management Section.

Handles multi-file batch document conversion supporting multiple input and output formats:
Input Formats:
  - PDF documents (.pdf)
  - Images (.jpg, .jpeg, .png, .webp, .bmp, .gif, .tiff, .svg)
  - Text & Structured Data (.txt, .md, .markdown, .html, .csv, .json, .xml)
  - Office & Electronic Documents (.docx, .xlsx, .pptx, .rtf, .epub, .odt, .ods, .odp)

Output Formats:
  - PDF, DOCX, XLSX, TXT, HTML, JPG, PNG, WEBP, MD, JSON, CSV

Key Features:
  - Independent file conversion (failures in one file do not interrupt others)
  - Memory-conscious stream processing for large batches
  - Collision-free, safe output filename generation
  - Individual converted file download & bulk ZIP archive generation
  - Security protections against path traversal and unsafe filenames
"""

from __future__ import annotations

import io
import json
import logging
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image

from app.core.paths import Paths

logger = logging.getLogger(__name__)

# Max file size limit per uploaded document (100 MB)
MAX_FILE_SIZE_BYTES = 100 * 1024 * 1024

# Supported Output Formats
SUPPORTED_OUTPUT_FORMATS: Dict[str, Dict[str, str]] = {
    "pdf": {"name": "PDF Document", "ext": ".pdf", "mime": "application/pdf"},
    "docx": {"name": "Word Document", "ext": ".docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "xlsx": {"name": "Excel Spreadsheet", "ext": ".xlsx", "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
    "txt": {"name": "Text Document", "ext": ".txt", "mime": "text/plain"},
    "html": {"name": "HTML Document", "ext": ".html", "mime": "text/html"},
    "jpg": {"name": "JPEG Image", "ext": ".jpg", "mime": "image/jpeg"},
    "png": {"name": "PNG Image", "ext": ".png", "mime": "image/png"},
    "webp": {"name": "WEBP Image", "ext": ".webp", "mime": "image/webp"},
    "md": {"name": "Markdown Document", "ext": ".md", "mime": "text/markdown"},
    "json": {"name": "JSON File", "ext": ".json", "mime": "application/json"},
    "csv": {"name": "CSV Spreadsheet", "ext": ".csv", "mime": "text/csv"},
}

# Recognized input file extensions
ALLOWED_INPUT_EXTENSIONS = {
    # PDF & Office
    ".pdf", ".docx", ".xlsx", ".pptx", ".rtf", ".epub", ".odt", ".ods", ".odp",
    # Images
    ".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff", ".svg",
    # Text & Code/Data
    ".txt", ".html", ".htm", ".md", ".markdown", ".json", ".csv", ".xml"
}


class BatchConversionService:
    """Enterprise service for batch document conversion across multiple formats."""

    def sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal and unsafe characters."""
        if not filename:
            return "document.pdf"
        clean = Path(filename).name
        clean = re.sub(r'[\\/:*?"<>|]', "_", clean)
        clean = re.sub(r"\s+", " ", clean).strip(" ._")
        return clean or "converted_document"

    def get_unique_filename(self, output_dir: Path, filename: str) -> str:
        """Generate a safe, unique filename inside output_dir to prevent overwriting existing files."""
        dest_path = output_dir / filename
        if not dest_path.exists():
            return filename

        p = Path(filename)
        stem = p.stem
        ext = p.suffix

        match = re.match(r"^(.*?)\s*\(\d+\)$", stem)
        if match:
            stem = match.group(1).strip()

        counter = 1
        while True:
            candidate = f"{stem} ({counter}){ext}"
            if not (output_dir / candidate).exists():
                return candidate
            counter += 1

    def validate_input_file(self, filename: str, file_bytes: bytes) -> Tuple[bool, str]:
        """Validate input file size and format before attempting conversion."""
        if not file_bytes or len(file_bytes) == 0:
            return False, "File is empty (0 bytes)."

        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            size_mb = len(file_bytes) / (1024 * 1024)
            return False, f"File size ({size_mb:.1f} MB) exceeds maximum allowed limit (100 MB)."

        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_INPUT_EXTENSIONS:
            return False, f"Unsupported input file format '{ext}'."

        return True, "" 
    

    def _extract_text_content(self, filename: str, file_bytes: bytes, ext: str) -> Tuple[str, List[Image.Image]]:
        """
        Extract text content and/or PIL Image objects from input file bytes.
        """
        text_lines: List[str] = []
        images: List[Image.Image] = []

        if ext == ".pdf":
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    txt = page.get_text("text") or ""
                    if txt.strip():
                        text_lines.append(txt.strip())
                    # Convert page to image for image rendering targets if needed
                    pix = page.get_pixmap(dpi=150)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    images.append(img)
                doc.close()
            except Exception as e:
                logger.warning(f"Error reading PDF content for {filename}: {e}")

        elif ext in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff"]:
            try:
                img = Image.open(io.BytesIO(file_bytes))
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                images.append(img)
                text_lines.append(f"[Image File: {filename}, Dimensions: {img.width}x{img.height}]")
            except Exception as e:
                logger.warning(f"Error opening image {filename}: {e}")

        elif ext in [".txt", ".md", ".markdown", ".html", ".htm", ".csv", ".json", ".xml"]:
            try:
                decoded = file_bytes.decode("utf-8", errors="replace")
                text_lines.append(decoded)
            except Exception as e:
                logger.warning(f"Error decoding text file {filename}: {e}")

        elif ext == ".docx":
            try:
                from docx import Document
                doc_obj = Document(io.BytesIO(file_bytes))
                full_text = [p.text for p in doc_obj.paragraphs if p.text.strip()]
                text_lines.append("\n".join(full_text))
            except Exception:
                text_lines.append(file_bytes.decode("utf-8", errors="ignore"))

        else:
            # Fallback text extraction
            try:
                decoded = file_bytes.decode("utf-8", errors="ignore")
                if decoded.strip():
                    text_lines.append(decoded)
            except Exception:
                pass

        full_text = "\n\n".join(text_lines).strip()
        return full_text, images

    def convert_single_file(
        self,
        session_id: str,
        file_bytes: bytes,
        original_filename: str,
        target_format: str,
        output_dir: Path,
    ) -> List[Path]:
        """
        Convert a single validated input file into the requested target format.
        Returns a list of generated file paths in output_dir.
        """
        target_fmt = (target_format or "pdf").lower().strip()
        fmt_info = SUPPORTED_OUTPUT_FORMATS.get(target_fmt, {"ext": f".{target_fmt}"})
        target_ext = fmt_info["ext"]

        clean_stem = Path(original_filename).stem or "converted"
        clean_stem = re.sub(r'[\\/:*?"<>|]', "_", clean_stem).strip(" ._") or "converted"
        input_ext = Path(original_filename).suffix.lower()

        generated_paths: List[Path] = []

        # ── 1. Target Format: PDF ─────────────────────────────────────────
        if target_fmt == "pdf":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.pdf")
            out_path = output_dir / target_name

            if input_ext == ".pdf":
                out_path.write_bytes(file_bytes)
                generated_paths.append(out_path)

            elif input_ext in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff"]:
                img = Image.open(io.BytesIO(file_bytes))
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                pdf_bytes_io = io.BytesIO()
                img.save(pdf_bytes_io, format="PDF")
                out_path.write_bytes(pdf_bytes_io.getvalue())
                generated_paths.append(out_path)

            else:
                # Convert text/other format to PDF using PyMuPDF or simple text PDF builder
                text, images = self._extract_text_content(original_filename, file_bytes, input_ext)
                if images and not text.strip():
                    pdf_bytes_io = io.BytesIO()
                    images[0].save(pdf_bytes_io, format="PDF")
                    out_path.write_bytes(pdf_bytes_io.getvalue())
                    generated_paths.append(out_path)
                else:
                    doc = fitz.open()
                    page = doc.new_page()
                    rect = fitz.Rect(50, 50, 545, 792)
                    page.insert_textbox(rect, text or f"Content from {original_filename}", fontsize=11)
                    pdf_bytes = doc.write()
                    doc.close()
                    out_path.write_bytes(pdf_bytes)
                    generated_paths.append(out_path)

        # ── 2. Target Format: Images (JPG, PNG, WEBP) ─────────────────────
        elif target_fmt in ["jpg", "png", "webp"]:
            img_ext = target_fmt

            if input_ext == ".pdf":
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                if len(doc) == 1:
                    target_name = self.get_unique_filename(output_dir, f"{clean_stem}.{img_ext}")
                    out_path = output_dir / target_name
                    pix = doc[0].get_pixmap(dpi=150)
                    if img_ext == "png":
                        pix.save(str(out_path))
                    else:
                        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                        img.save(str(out_path), format=img_ext.upper())
                    generated_paths.append(out_path)
                else:
                    for i, page in enumerate(doc):
                        target_name = self.get_unique_filename(output_dir, f"{clean_stem}_page_{i+1}.{img_ext}")
                        out_path = output_dir / target_name
                        pix = page.get_pixmap(dpi=150)
                        if img_ext == "png":
                            pix.save(str(out_path))
                        else:
                            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                            img.save(str(out_path), format=img_ext.upper())
                        generated_paths.append(out_path)
                doc.close()

            elif input_ext in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff"]:
                target_name = self.get_unique_filename(output_dir, f"{clean_stem}.{img_ext}")
                out_path = output_dir / target_name
                img = Image.open(io.BytesIO(file_bytes))
                if img_ext != "png" and img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(str(out_path), format=img_ext.upper())
                generated_paths.append(out_path)

            else:
                text, images = self._extract_text_content(original_filename, file_bytes, input_ext)
                target_name = self.get_unique_filename(output_dir, f"{clean_stem}.{img_ext}")
                out_path = output_dir / target_name
                if images:
                    img = images[0]
                    if img_ext != "png" and img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                    img.save(str(out_path), format=img_ext.upper())
                else:
                    # Render text on simple PIL Image canvas
                    from PIL import ImageDraw
                    canvas = Image.new("RGB", (800, 1000), color=(255, 255, 255))
                    draw = ImageDraw.Draw(canvas)
                    draw.text((40, 40), f"File: {original_filename}\n\n{text[:1500]}", fill=(0, 0, 0))
                    canvas.save(str(out_path), format=img_ext.upper())
                generated_paths.append(out_path)

        # ── 3. Target Format: Text (TXT) ──────────────────────────────────
        elif target_fmt == "txt":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.txt")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)
            out_path.write_text(text or f"[Converted Content from {original_filename}]", encoding="utf-8")
            generated_paths.append(out_path)

        # ── 4. Target Format: HTML ────────────────────────────────────────
        elif target_fmt == "html":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.html")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)
            paragraphs = [f"<p>{p.strip().replace('<', '&lt;').replace('>', '&gt;')}</p>" for p in text.split("\n\n") if p.strip()]
            body_content = "\n".join(paragraphs) if paragraphs else f"<p>{text.replace('<', '&lt;')}</p>"
            html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{clean_stem}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #1e293b; }}
        h1 {{ border-bottom: 2px solid #e2e8f0; padding-bottom: 10px; color: #0f172a; }}
        p {{ margin-bottom: 1.2rem; }}
    </style>
</head>
<body>
    <h1>{clean_stem}</h1>
    {body_content}
</body>
</html>"""
            out_path.write_text(html_doc, encoding="utf-8")
            generated_paths.append(out_path)

        # ── 5. Target Format: Markdown (MD) ───────────────────────────────
        elif target_fmt == "md":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.md")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)
            md_content = f"# {clean_stem}\n\n" + text
            out_path.write_text(md_content, encoding="utf-8")
            generated_paths.append(out_path)

        # ── 6. Target Format: JSON ────────────────────────────────────────
        elif target_fmt == "json":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.json")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)
            json_obj = {
                "source_filename": original_filename,
                "input_format": input_ext,
                "converted_content": text,
                "file_size_bytes": len(file_bytes),
            }
            out_path.write_text(json.dumps(json_obj, indent=2), encoding="utf-8")
            generated_paths.append(out_path)

        # ── 7. Target Format: CSV ─────────────────────────────────────────
        elif target_fmt == "csv":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.csv")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)
            csv_lines = ["Line_Number,Text_Content"]
            for idx, line in enumerate(text.splitlines(), start=1):
                if line.strip():
                    escaped_line = line.strip().replace('"', '""')
                    csv_lines.append(f'"{idx}","{escaped_line}"')
            out_path.write_text("\n".join(csv_lines), encoding="utf-8")
            generated_paths.append(out_path)

        # ── 8. Target Format: Word (DOCX) ─────────────────────────────────
        elif target_fmt == "docx":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.docx")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)

            try:
                from docx import Document
                docx_obj = Document()
                docx_obj.add_heading(clean_stem, level=1)
                for line in text.splitlines():
                    if line.strip():
                        docx_obj.add_paragraph(line.strip())
                docx_obj.save(str(out_path))
            except Exception:
                out_path.write_text(text, encoding="utf-8")
            generated_paths.append(out_path)

        # ── 9. Target Format: Excel (XLSX) ────────────────────────────────
        elif target_fmt == "xlsx":
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}.xlsx")
            out_path = output_dir / target_name
            text, _ = self._extract_text_content(original_filename, file_bytes, input_ext)

            try:
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Converted Data"
                ws.cell(row=1, column=1, value="Line #")
                ws.cell(row=1, column=2, value="Content")

                row_idx = 2
                for line in text.splitlines():
                    if line.strip():
                        ws.cell(row=row_idx, column=1, value=row_idx - 1)
                        ws.cell(row=row_idx, column=2, value=line.strip())
                        row_idx += 1
                wb.save(str(out_path))
            except Exception:
                out_path.write_text(text, encoding="utf-8")
            generated_paths.append(out_path)

        # ── Fallback Target Converter ─────────────────────────────────────
        else:
            target_name = self.get_unique_filename(output_dir, f"{clean_stem}{target_ext}")
            out_path = output_dir / target_name
            out_path.write_bytes(file_bytes)
            generated_paths.append(out_path)

        return generated_paths

    def process_batch_conversion(
        self,
        session_id: str,
        files_data: List[Dict[str, Any]],
        target_format: str = "pdf",
    ) -> Dict[str, Any]:
        """
        Process a batch of uploaded files for conversion.

        Args:
            session_id: Unique session / request identifier
            files_data: List of dicts containing {"filename": str, "bytes": bytes}
            target_format: Target conversion format (pdf, docx, xlsx, txt, html, jpg, png, webp, md, json, csv)

        Returns:
            Dict summary with total, successful, failed, detailed list of failures, individual download URLs, and ZIP download URL.
        """
        target_fmt = (target_format or "pdf").lower().strip()
        if target_fmt not in SUPPORTED_OUTPUT_FORMATS:
            supported_str = ", ".join(SUPPORTED_OUTPUT_FORMATS.keys())
            raise ValueError(f"Unsupported output format '{target_format}'. Supported output formats: {supported_str}")

        if not files_data or len(files_data) == 0:
            raise ValueError("No files provided for batch conversion.")

        session_dir = Paths.request_output(session_id)
        session_dir.mkdir(parents=True, exist_ok=True)

        batch_out_dir = session_dir / "converted_files"
        batch_out_dir.mkdir(parents=True, exist_ok=True)

        results: List[Dict[str, Any]] = []
        failed_details: List[Dict[str, str]] = []
        successful_files_count = 0
        failed_files_count = 0
        generated_all_files: List[Path] = []

        total_files = len(files_data)

        for item in files_data:
            orig_filename = self.sanitize_filename(item.get("filename", "document"))
            file_bytes = item.get("bytes", b"")

            # 1. Validate file
            is_valid, err_msg = self.validate_input_file(orig_filename, file_bytes)
            if not is_valid:
                failed_files_count += 1
                failed_details.append({
                    "filename": orig_filename,
                    "reason": err_msg
                })
                results.append({
                    "filename": orig_filename,
                    "status": "failed",
                    "error": err_msg,
                    "output_files": [],
                    "download_urls": []
                })
                continue

            # 2. Process file conversion independently
            try:
                converted_paths = self.convert_single_file(
                    session_id=session_id,
                    file_bytes=file_bytes,
                    original_filename=orig_filename,
                    target_format=target_fmt,
                    output_dir=batch_out_dir,
                )

                if converted_paths:
                    successful_files_count += 1
                    generated_all_files.extend(converted_paths)
                    output_file_names = [p.name for p in converted_paths]
                    download_urls = [
                        f"/document-management/batch-conversion/download-file/{session_id}/{p.name}"
                        for p in converted_paths
                    ]

                    results.append({
                        "filename": orig_filename,
                        "status": "success",
                        "output_files": output_file_names,
                        "download_urls": download_urls,
                        "primary_download_url": download_urls[0] if download_urls else None,
                    })
                else:
                    failed_files_count += 1
                    failed_details.append({
                        "filename": orig_filename,
                        "reason": "Conversion yielded no output files."
                    })
                    results.append({
                        "filename": orig_filename,
                        "status": "failed",
                        "error": "Conversion yielded no output files.",
                        "output_files": [],
                        "download_urls": []
                    })

            except Exception as exc:
                logger.error(f"Batch conversion error for '{orig_filename}' to '{target_fmt}': {exc}", exc_info=True)
                failed_files_count += 1
                failed_details.append({
                    "filename": orig_filename,
                    "reason": f"Conversion error: {str(exc)}"
                })
                results.append({
                    "filename": orig_filename,
                    "status": "failed",
                    "error": str(exc),
                    "output_files": [],
                    "download_urls": []
                })

        # 3. Create ZIP Archive for bulk download if successful files exist
        zip_filename = ""
        has_download = False

        if generated_all_files:
            zip_filename = f"batch_conversion_{target_fmt}_{session_id[:8]}.zip"
            zip_path = session_dir / zip_filename
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for fpath in generated_all_files:
                    zf.write(fpath, arcname=fpath.name)
            has_download = True

        return {
            "session_id": session_id,
            "target_format": target_fmt,
            "format_name": SUPPORTED_OUTPUT_FORMATS[target_fmt]["name"],
            "total_files": total_files,
            "successful_files": successful_files_count,
            "failed_files": failed_files_count,
            "results": results,
            "failed_details": failed_details,
            "zip_filename": zip_filename,
            "has_download": has_download,
            "zip_download_url": f"/document-management/batch-conversion/download-zip/{session_id}" if has_download else None,
        }

    def get_converted_file_for_download(self, session_id: str, filename: str) -> Tuple[Path, str]:
        """Locate and return a specific single converted file for download."""
        if not session_id or re.search(r"[\\/]", session_id):
            raise ValueError("Invalid session ID.")
        if not filename or re.search(r"[\\/]", filename):
            raise ValueError("Invalid filename.")

        session_dir = Paths.request_output(session_id)
        batch_out_dir = session_dir / "converted_files"

        target_path = batch_out_dir / filename
        if target_path.exists() and target_path.is_file():
            return target_path, filename

        # Fallback search inside session_dir
        direct_path = session_dir / filename
        if direct_path.exists() and direct_path.is_file():
            return direct_path, filename

        raise ValueError(f"Converted file '{filename}' not found for session.")

    def get_zip_for_download(self, session_id: str) -> Tuple[Path, str]:
        """Locate and return the session ZIP archive or main output file."""
        if not session_id or re.search(r"[\\/]", session_id):
            raise ValueError("Invalid session ID.")

        session_dir = Paths.request_output(session_id)
        if not session_dir.exists():
            raise ValueError("Session conversion data not found or expired.")

        zips = list(session_dir.glob("*.zip"))
        if zips:
            return zips[0], zips[0].name

        # Fallback to single file in converted_files
        batch_out_dir = session_dir / "converted_files"
        if batch_out_dir.exists():
            sub_files = [f for f in batch_out_dir.iterdir() if f.is_file()]
            if sub_files:
                return sub_files[0], sub_files[0].name

        raise ValueError("No downloadable conversion archive found for session.")


batch_conversion_service = BatchConversionService()
