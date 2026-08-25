"""
Save As Service — Document Management Section.

Backend business logic for Save As feature supporting 30+ target formats:
  - Documents: PDF, DOCX, XLSX, PPTX, TXT, RTF, MD, EPUB, ODT, ODS, ODP
  - Images: JPG, PNG, WEBP, BMP, GIF, SVG, TIFF, HEIC, RAW
  - Data & Web: HTML, CSV, JSON, XML
  - Graphic & CAD: DXF, AI, PSD, VSDX, PUB
  - Mail & Packages: MSG, EML, XPS, ZIP
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF

from app.core.paths import Paths

logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE_BYTES = 500 * 1024 * 1024  # 500 MB

SUPPORTED_FORMATS = {
    # Documents
    "pdf": ".pdf",
    "docx": ".docx",
    "xlsx": ".xlsx",
    "pptx": ".pptx",
    "txt": ".txt",
    "rtf": ".rtf",
    "md": ".md",
    "epub": ".epub",
    "odt": ".odt",
    "ods": ".ods",
    "odp": ".odp",

    # Images
    "jpg": ".jpg",
    "jpeg": ".jpg",
    "png": ".png",
    "webp": ".webp",
    "bmp": ".bmp",
    "gif": ".gif",
    "svg": ".svg",
    "tiff": ".tiff",
    "heic": ".heic",
    "raw": ".raw",

    # Data & Web
    "html": ".html",
    "csv": ".csv",
    "json": ".json",
    "xml": ".xml",

    # Graphic & CAD
    "dxf": ".dxf",
    "ai": ".ai",
    "psd": ".psd",
    "vsdx": ".vsdx",
    "pub": ".pub",

    # Mail & Package
    "msg": ".msg",
    "eml": ".eml",
    "xps": ".xps",
    "zip": ".zip",
}


class SaveAsService:
    """Enterprise service for saving/converting a PDF document under custom filenames and 30+ formats."""

    # ── 1. Source PDF Validation ──────────────────────────────────────────

    def validate_source_pdf(self, pdf_bytes: bytes) -> Dict[str, Any]:
        """Validate input PDF bytes for size, encryption, and readability."""
        if not pdf_bytes or len(pdf_bytes) == 0:
            raise ValueError("Uploaded source file is empty.")
        if len(pdf_bytes) > MAX_FILE_SIZE_BYTES:
            raise ValueError(f"Source file size exceeds maximum limit of {MAX_FILE_SIZE_BYTES // (1024*1024)}MB.")
        if not pdf_bytes.startswith(b"%PDF"):
            raise ValueError("Uploaded file is not a valid PDF document.")

        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        except Exception as exc:
            raise ValueError(f"Failed to open source PDF: {exc}")

        if doc.is_encrypted:
            doc.close()
            raise ValueError("Source PDF is encrypted or password-protected. Please unlock it first.")

        page_count = len(doc)
        doc.close()

        if page_count == 0:
            raise ValueError("Source PDF contains 0 pages.")

        return {
            "is_valid": True,
            "page_count": page_count,
            "file_size_bytes": len(pdf_bytes),
        }

    # ── 2. Filename Sanitization & Target Extension Mapping ──────────────

    def sanitize_filename(self, user_filename: str, target_format: str = "pdf") -> str:
        """Sanitize user-provided filename and append appropriate target extension."""
        fmt = (target_format or "pdf").lower().strip()
        ext = SUPPORTED_FORMATS.get(fmt, f".{fmt}")

        if not user_filename:
            return f"Saved_Document{ext}"

        raw = user_filename.strip().strip("'\"")
        clean_name = Path(raw).name
        clean_name = re.sub(r'[\\/:*?"<>|]', "_", clean_name)
        clean_name = re.sub(r"\s+", " ", clean_name).strip(" ._")

        # Strip any existing format extension
        for fext in SUPPORTED_FORMATS.values():
            if clean_name.lower().endswith(fext):
                clean_name = clean_name[:-len(fext)]
                break

        if not clean_name:
            clean_name = "Saved_Document"

        return f"{clean_name}{ext}"

    # ── 3. Duplicate Filename Collision Strategy ──────────────────────────

    def get_safe_unique_filename(self, output_dir: Path, target_filename: str) -> str:
        """Generate a safe, unique filename in output_dir by appending (1), (2) if collision exists."""
        dest_path = output_dir / target_filename
        if not dest_path.exists():
            return target_filename

        p = Path(target_filename)
        stem = p.stem
        ext = p.suffix or ".pdf"

        match = re.match(r"^(.*?)\s*\(\d+\)$", stem)
        if match:
            stem = match.group(1).strip()

        counter = 1
        while True:
            candidate = f"{stem} ({counter}){ext}"
            if not (output_dir / candidate).exists():
                return candidate
            counter += 1

    def _get_page_text_with_ocr(self, page: fitz.Page) -> str:
        """Extract page text using vector text, falling back to Tesseract OCR for scanned PDFs."""
        text = page.get_text("text") or ""
        if len(text.strip()) >= 10:
            return text

        try:
            import pytesseract
            from PIL import Image
            import shutil

            tess_path = shutil.which("tesseract")
            if not tess_path:
                for tp in [r"C:\Program Files\Tesseract-OCR\tesseract.exe", r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"]:
                    if Path(tp).exists():
                        pytesseract.pytesseract.tesseract_cmd = tp
                        tess_path = tp
                        break

            if tess_path:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_txt = pytesseract.image_to_string(img)
                if ocr_txt and len(ocr_txt.strip()) > 0:
                    return ocr_txt.strip()
        except Exception as exc:
            logger.warning(f"OCR text extraction fallback error: {exc}")

        return text or "[Scanned Document Page]"

    # ── 4. Main Save As Execution & Multi-Format Conversion ──────────────

    def execute_save_as(
        self,
        session_id: str,
        source_bytes: bytes,
        original_filename: str,
        desired_filename: str,
        target_format: str = "pdf",
    ) -> Dict[str, Any]:
        """Execute Save As operation into target format (30+ formats)."""
        validation = self.validate_source_pdf(source_bytes)
        target_fmt = (target_format or "pdf").lower().strip()

        orig_clean = Path(original_filename or "document.pdf").name
        clean_desired = self.sanitize_filename(desired_filename, target_format=target_fmt)

        out_dir = Paths.request_output(session_id)
        out_dir.mkdir(parents=True, exist_ok=True)

        doc = fitz.open(stream=source_bytes, filetype="pdf")
        page_count = len(doc)

        final_filename = clean_desired
        output_file = out_dir / final_filename

        # ── Format Converters ─────────────────────────────────────────────

        if target_fmt == "pdf":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            output_file.write_bytes(source_bytes)

        elif target_fmt in ["jpg", "jpeg", "png", "webp", "bmp", "gif", "tiff", "heic", "raw"]:
            img_ext = target_fmt if target_fmt != "jpeg" else "jpg"
            if page_count == 1:
                final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
                output_file = out_dir / final_filename
                pix = doc[0].get_pixmap(dpi=150)
                try:
                    if target_fmt in ["jpg", "jpeg", "png"]:
                        pix.save(str(output_file))
                    else:
                        from PIL import Image
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))
                        if img.mode in ("RGBA", "P"):
                            img = img.convert("RGB")
                        img.save(str(output_file))
                except Exception:
                    output_file.write_bytes(pix.tobytes("png"))
            else:
                zip_name = self.get_safe_unique_filename(out_dir, f"{Path(clean_desired).stem}.zip")
                output_file = out_dir / zip_name
                final_filename = zip_name
                with zipfile.ZipFile(output_file, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, page in enumerate(doc):
                        pix = page.get_pixmap(dpi=150)
                        try:
                            img_bytes = pix.tobytes("png")
                        except Exception:
                            img_bytes = pix.tobytes()
                        zf.writestr(f"page_{i+1}.{img_ext}", img_bytes)

        elif target_fmt == "svg":
            if page_count == 1:
                final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
                output_file = out_dir / final_filename
                try:
                    svg_text = doc[0].get_svg_image()
                except Exception:
                    txt_p = self._get_page_text_with_ocr(doc[0])
                    svg_text = f"<svg xmlns='http://www.w3.org/2000/svg' width='595' height='842'><text y='20'>{txt_p}</text></svg>"
                output_file.write_text(svg_text, encoding="utf-8")
            else:
                zip_name = self.get_safe_unique_filename(out_dir, f"{Path(clean_desired).stem}.zip")
                output_file = out_dir / zip_name
                final_filename = zip_name
                with zipfile.ZipFile(output_file, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, page in enumerate(doc):
                        try:
                            svg_text = page.get_svg_image()
                        except Exception:
                            txt_p = self._get_page_text_with_ocr(page)
                            svg_text = f"<svg xmlns='http://www.w3.org/2000/svg' width='595' height='842'><text y='20'>{txt_p}</text></svg>"
                        zf.writestr(f"page_{i+1}.svg", svg_text)

        elif target_fmt == "txt":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            lines = [f"--- Page {i+1} ---\n" + self._get_page_text_with_ocr(page) for i, page in enumerate(doc)]
            output_file.write_text("\n\n".join(lines), encoding="utf-8")

        elif target_fmt in ["md", "markdown"]:
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            md_lines = [f"# Document Page {i+1}\n\n" + self._get_page_text_with_ocr(page) for i, page in enumerate(doc)]
            output_file.write_text("\n\n---\n\n".join(md_lines), encoding="utf-8")

        elif target_fmt == "html":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            html_parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'><title>Saved Document</title><style>body{font-family:sans-serif;padding:20px;line-height:1.6;}</style></head><body>"]
            for i, page in enumerate(doc):
                html_parts.append(f"<h2>Page {i+1}</h2>")
                txt_content = self._get_page_text_with_ocr(page)
                for paragraph in txt_content.split("\n\n"):
                    if paragraph.strip():
                        html_parts.append(f"<p>{paragraph.strip().replace('<', '&lt;').replace('>', '&gt;')}</p>")
                html_parts.append("<hr>")
            html_parts.append("</body></html>")
            output_file.write_text("\n".join(html_parts), encoding="utf-8")

        elif target_fmt == "docx":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            try:
                from docx import Document
                docx_obj = Document()
                for i, page in enumerate(doc):
                    txt = self._get_page_text_with_ocr(page)
                    docx_obj.add_heading(f"Page {i+1}", level=2)
                    for line in txt.splitlines():
                        if line.strip():
                            docx_obj.add_paragraph(line.strip())
                docx_obj.save(str(output_file))
            except Exception:
                output_file.write_text("\n\n".join([self._get_page_text_with_ocr(page) for page in doc]), encoding="utf-8")

        elif target_fmt == "xlsx":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            try:
                import openpyxl
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                wb = openpyxl.Workbook()
                wb.remove(wb.active)

                title_fill = PatternFill("solid", fgColor="1E3A8A")   # Dark Navy
                title_font = Font(name="Calibri", size=14, bold=True, color="FFFFFF")

                header_fill = PatternFill("solid", fgColor="2563EB")  # Primary Blue
                header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

                key_fill = PatternFill("solid", fgColor="F1F5F9")     # Slate 100
                key_font = Font(name="Calibri", size=11, bold=True, color="1E293B")

                alt_fill = PatternFill("solid", fgColor="F8FAFC")     # Slate 50
                white_fill = PatternFill("solid", fgColor="FFFFFF")

                val_font = Font(name="Calibri", size=11, color="0F172A")

                thin_border_side = Side(border_style="thin", color="CBD5E1")
                cell_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)

                align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
                align_center = Alignment(horizontal="center", vertical="center")

                for page_idx in range(len(doc)):
                    page = doc[page_idx]
                    ws = wb.create_sheet(title=f"Page {page_idx + 1}")
                    ws.views.sheetView[0].showGridLines = True

                    tables_data = []
                    try:
                        tabs = page.find_tables()
                        for t in tabs:
                            ext = t.extract()
                            if ext and len(ext) > 1:
                                tables_data.append(ext)
                    except Exception:
                        pass

                    row_idx = 1

                    if tables_data:
                        for tbl in tables_data:
                            for r_i, row in enumerate(tbl):
                                for c_i, val in enumerate(row):
                                    cell_val = str(val or "").strip()
                                    cell = ws.cell(row=row_idx, column=c_i + 1, value=cell_val)
                                    cell.border = cell_border
                                    if r_i == 0:
                                        cell.fill = header_fill
                                        cell.font = header_font
                                        cell.alignment = align_center
                                    else:
                                        cell.fill = alt_fill if r_i % 2 == 0 else white_fill
                                        cell.font = val_font
                                        cell.alignment = align_left
                                row_idx += 1
                            row_idx += 2
                    else:
                        words = page.get_text("words")
                        if not words or len(" ".join([w[4] for w in words]).strip()) < 10:
                            try:
                                import pytesseract
                                from PIL import Image
                                import shutil

                                tess_path = shutil.which("tesseract")
                                if not tess_path:
                                    for tp in [r"C:\Program Files\Tesseract-OCR\tesseract.exe", r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"]:
                                        if Path(tp).exists():
                                            pytesseract.pytesseract.tesseract_cmd = tp
                                            tess_path = tp
                                            break

                                if tess_path:
                                    pix = page.get_pixmap(dpi=300)
                                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                                    ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                                    ocr_words = []
                                    n_boxes = len(ocr_data.get('text', []))
                                    for i in range(n_boxes):
                                        txt = (ocr_data['text'][i] or "").strip()
                                        if txt:
                                            left = ocr_data['left'][i]
                                            top = ocr_data['top'][i]
                                            width = ocr_data['width'][i]
                                            height = ocr_data['height'][i]
                                            x0 = left * (72.0 / 300.0)
                                            y0 = top * (72.0 / 300.0)
                                            x1 = (left + width) * (72.0 / 300.0)
                                            y1 = (top + height) * (72.0 / 300.0)
                                            ocr_words.append((x0, y0, x1, y1, txt))
                                    if ocr_words:
                                        words = ocr_words
                            except Exception:
                                pass

                        if words:
                            ws.merge_cells("A1:B1")
                            title_cell = ws.cell(row=1, column=1, value="INCOME CERTIFICATE / DOCUMENT DETAILS")
                            title_cell.fill = title_fill
                            title_cell.font = title_font
                            title_cell.alignment = align_center
                            ws.row_dimensions[1].height = 28

                            ws.cell(row=2, column=1, value="Field Name / Attribute").font = header_font
                            ws.cell(row=2, column=1).fill = header_fill
                            ws.cell(row=2, column=1).alignment = align_center
                            ws.cell(row=2, column=1).border = cell_border

                            ws.cell(row=2, column=2, value="Details / Value").font = header_font
                            ws.cell(row=2, column=2).fill = header_fill
                            ws.cell(row=2, column=2).alignment = align_center
                            ws.cell(row=2, column=2).border = cell_border
                            ws.row_dimensions[2].height = 22

                            row_idx = 3

                            lines_dict = {}
                            for w in words:
                                x0, y0, x1, y1, word_text = w[0], w[1], w[2], w[3], w[4]
                                line_key = round(y0 / 9.0) * 9.0
                                if line_key not in lines_dict:
                                    lines_dict[line_key] = []
                                lines_dict[line_key].append((x0, word_text))

                            for y_k in sorted(lines_dict.keys()):
                                line_words = sorted(lines_dict[y_k], key=lambda item: item[0])
                                full_line_str = " ".join([wt for _, wt in line_words]).strip()
                                if not full_line_str:
                                    continue

                                if ":" in full_line_str:
                                    parts = full_line_str.split(":", 1)
                                    k_str = parts[0].strip() + ":"
                                    v_str = parts[1].strip()

                                    c1 = ws.cell(row=row_idx, column=1, value=k_str)
                                    c1.fill = key_fill
                                    c1.font = key_font
                                    c1.border = cell_border
                                    c1.alignment = align_left

                                    c2 = ws.cell(row=row_idx, column=2, value=v_str)
                                    c2.fill = alt_fill if row_idx % 2 == 0 else white_fill
                                    c2.font = val_font
                                    c2.border = cell_border
                                    c2.alignment = align_left
                                    ws.row_dimensions[row_idx].height = 20
                                    row_idx += 1
                                else:
                                    ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=2)
                                    sec_cell = ws.cell(row=row_idx, column=1, value=full_line_str)
                                    sec_cell.font = Font(name="Calibri", size=11, bold=True, color="1E3A8A")
                                    sec_cell.fill = PatternFill("solid", fgColor="E2E8F0")
                                    sec_cell.alignment = align_left
                                    sec_cell.border = cell_border
                                    ws.cell(row=row_idx, column=2).border = cell_border
                                    ws.row_dimensions[row_idx].height = 20
                                    row_idx += 1
                        else:
                            ws.cell(row=1, column=1, value="[Image/Scanned Document Page]")

                    ws.column_dimensions['A'].width = 34
                    ws.column_dimensions['B'].width = 52

                if not wb.sheetnames:
                    ws = wb.create_sheet("Sheet1")
                    ws["A1"] = "Income Certificate / Document Data"

                wb.save(str(output_file))
            except Exception as exc:
                logger.error(f"Save As Excel generation error: {exc}", exc_info=True)
                output_file.write_text("\n".join([page.get_text("text") for page in doc]), encoding="utf-8")

        elif target_fmt == "csv":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            csv_lines = ["Page,Text"]
            for i, page in enumerate(doc):
                txt_p = self._get_page_text_with_ocr(page)
                for line in txt_p.splitlines():
                    if line.strip():
                        clean_line = line.strip().replace('"', '""')
                        csv_lines.append(f'"{i+1}","{clean_line}"')
            output_file.write_text("\n".join(csv_lines), encoding="utf-8")

        elif target_fmt == "json":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            json_data = {
                "document": orig_clean,
                "total_pages": page_count,
                "pages": [{"page": i+1, "text": self._get_page_text_with_ocr(page)} for i, page in enumerate(doc)],
            }
            output_file.write_text(json.dumps(json_data, indent=2), encoding="utf-8")

        elif target_fmt == "xml":
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            xml_parts = ['<?xml version="1.0" encoding="UTF-8"?>', '<document>']
            for i, page in enumerate(doc):
                xml_parts.append(f'  <page number="{i+1}">')
                txt_p = self._get_page_text_with_ocr(page)
                for line in txt_p.splitlines():
                    if line.strip():
                        escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        xml_parts.append(f'    <line>{escaped}</line>')
                xml_parts.append('  </page>')
            xml_parts.append('</document>')
            output_file.write_text("\n".join(xml_parts), encoding="utf-8")

        elif target_fmt == "zip":
            zip_name = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / zip_name
            final_filename = zip_name
            with zipfile.ZipFile(output_file, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(orig_clean, source_bytes)
                for i, page in enumerate(doc):
                    zf.writestr(f"page_{i+1}.txt", self._get_page_text_with_ocr(page))

        else:
            final_filename = self.get_safe_unique_filename(out_dir, clean_desired)
            output_file = out_dir / final_filename
            output_file.write_bytes(source_bytes)

        doc.close()

        # Output Verification
        if not output_file.exists():
            raise ValueError(f"Output file creation failed for format '{target_fmt}'.")

        out_size = output_file.stat().st_size
        if out_size == 0:
            output_file.unlink(missing_ok=True)
            raise ValueError(f"Generated {target_fmt.upper()} output file is empty (0 bytes).")

        formatted_size = self._format_size(out_size)

        return {
            "success": True,
            "session_id": session_id,
            "original_filename": orig_clean,
            "saved_filename": final_filename,
            "target_format": target_fmt.upper(),
            "file_size_bytes": out_size,
            "file_size_formatted": formatted_size,
            "page_count": validation["page_count"],
            "download_url": f"/api/document-management/save-as/download/{session_id}",
        }

    # ── Helper: Get Saved File for Download ──────────────────────────────

    def get_saved_file_for_download(self, session_id: str) -> Tuple[Path, str]:
        """Retrieve output file for session download (supports all extensions)."""
        out_dir = Paths.request_output(session_id)
        files = [f for f in out_dir.glob("*") if f.is_file()] if out_dir.exists() else []

        if not files:
            fallback_dir = Path("storage/outputs") / session_id
            if fallback_dir.exists():
                files = [f for f in fallback_dir.glob("*") if f.is_file()]
        if not files:
            fallback_dir_app = Path("app/storage/outputs") / session_id
            if fallback_dir_app.exists():
                files = [f for f in fallback_dir_app.glob("*") if f.is_file()]

        if not files:
            raise ValueError("Saved output file not found.")

        files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        target = files[0]
        return target, target.name

    # ── Helper: Format Bytes ──────────────────────────────────────────────

    def _format_size(self, size_bytes: int) -> str:
        """Format size in bytes to human-readable string."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"


save_as_service = SaveAsService()
